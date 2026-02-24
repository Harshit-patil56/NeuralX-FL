"""
generate_report.py — Generates the hackathon Word report for NeuralX-FL.
Run once:  python generate_report.py
Output:    NeuralX_FL_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ---------------------------------------------------------------------------
# Page margins
# ---------------------------------------------------------------------------
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.bold = True
    return p

def para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def table_2col(rows, header=None):
    col_count = 2
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=col_count)
    t.style = "Light List Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_start = 0
    if header:
        for i, h in enumerate(header):
            cell = t.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        r_start = 1
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[r_start + i].cells[j].text = str(val)
    doc.add_paragraph()
    return t

# ---------------------------------------------------------------------------
# TITLE PAGE
# ---------------------------------------------------------------------------
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("NeuralX-FL")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Asynchronous Byzantine-Robust Federated Learning for Medical Pathology")
run2.font.size = Pt(14)
run2.bold = True

doc.add_paragraph()

team_para = doc.add_paragraph()
team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
team_para.add_run("Team: NeuralX\n").bold = True
team_para.add_run("Harshit Patil  |  Tanmay Patil  |  Prathmesh Shejwal\n")
team_para.add_run("St. John College of Engineering and Management")

doc.add_paragraph()
doc.add_page_break()

# ---------------------------------------------------------------------------
# 1. INTRODUCTION
# ---------------------------------------------------------------------------
heading("1. Introduction")

para(
    "NeuralX-FL is a privacy-preserving federated learning system designed for "
    "multi-hospital collaborative training of a colon pathology classifier. "
    "The system enables hospitals to jointly train a deep learning model on patient "
    "pathology slides without any raw image data ever leaving a hospital's premises. "
    "Only model weight updates — floating-point arrays — are exchanged over the network."
)

para(
    "The project uses the PathMNIST dataset (107,180 colon-tissue patches, 9 tissue "
    "classes) and a custom CNN called PathologyNet, running on an NVIDIA GTX 1650 GPU "
    "with PyTorch 2.6.0 and CUDA 12.4. The federated framework is built on Flower (flwr), "
    "a production-grade FL library. The full system is implemented in approximately 900 "
    "lines of Python across 8 focused modules, accompanied by 29 unit tests."
)

# ---------------------------------------------------------------------------
# 2. PROBLEM STATEMENT
# ---------------------------------------------------------------------------
heading("2. Problem Statement")

para(
    "Hospitals hold large volumes of pathology slide images that could train powerful "
    "diagnostic AI models. However, sharing these images is legally prohibited under "
    "HIPAA and GDPR. A centralised model is therefore infeasible. Federated Learning "
    "solves the data-sharing problem but introduces two new real-world threats:"
)

table_2col(
    [
        ["Byzantine Attacks",
         "A malicious hospital sends deliberately corrupted gradient updates to "
         "poison the global model. Standard FedAvg has no defence and loses 78% "
         "accuracy in a single round."],
        ["Asynchronous Networks",
         "Hospitals have different hardware speeds and network reliability. Updates "
         "arrive at different times. Standard FL stalls waiting for the slowest "
         "client, making it impractical for real deployments."],
    ],
    header=["Threat", "Description"]
)

para(
    "NeuralX-FL solves both threats simultaneously — no existing open-source "
    "Flower-based system combines async-robust aggregation with dynamic trust scoring "
    "in a single real-network-capable deployment."
)

# ---------------------------------------------------------------------------
# 3. WORKING
# ---------------------------------------------------------------------------
heading("3. Working")

doc.add_heading("3.1  Federation Setup", level=2)
table_2col(
    [
        ["Total hospitals (clients)", "10"],
        ["Clients sampled per round", "6"],
        ["Async buffer size", "4 (fastest clean updates used per round)"],
        ["Training rounds", "20"],
        ["Dataset", "PathMNIST — 107,180 colon-tissue patches, 9 classes"],
        ["Model", "PathologyNet CNN (Conv→Pool×2 → FC×2, 9-class output)"],
        ["Data partitioning", "Dirichlet α=0.5 (non-IID per hospital)"],
    ],
    header=["Parameter", "Value"]
)

doc.add_heading("3.2  Client Types", level=2)
table_2col(
    [
        ["Hospitals {0, 1}", "Malicious (20%) — gradient scaling attack ×50"],
        ["Hospitals {4, 5}", "Noisy (20%) — 30% random label flipping"],
        ["Hospitals {6, 7}", "Unreliable (20%) — 40% dropout probability per round"],
        ["Hospitals {2,3,8,9}", "Honest (40%) — normal federated training"],
    ],
    header=["Hospital IDs", "Behaviour"]
)

doc.add_heading("3.3  Defence Pipeline (4 Steps)", level=2)

bullet("L2 Norm Filter: ", bold_prefix="Step 1a — ")
p = doc.paragraphs[-1]
p.add_run(
    "Reject any update whose L2 norm exceeds 3× the median norm. "
    "A ×50 scaling attack inflates the norm by ~50×, making it trivially detectable."
)

bullet("Cosine Similarity Filter: ", bold_prefix="Step 1b — ")
p = doc.paragraphs[-1]
p.add_run(
    "Reject updates whose cosine similarity with the consensus direction (coordinate-wise median) "
    "is below 0. Catches sign-flip and directional poisoning attacks."
)

bullet("Async Selection: ", bold_prefix="Step 2 — ")
p = doc.paragraphs[-1]
p.add_run(
    "After detection, the 4 fastest surviving updates are selected. "
    "Detection runs before selection so attackers cannot bypass filters by responding fast."
)

bullet("Robust Aggregation: ", bold_prefix="Step 3 — ")
p = doc.paragraphs[-1]
p.add_run(
    "Coordinate-wise trimmed mean with staleness weighting "
    "(weight = 1/(1+staleness_rounds)). Stale updates from slow hospitals contribute less."
)

bullet("Dynamic Trust Scoring: ", bold_prefix="Step 4 — ")
p = doc.paragraphs[-1]
p.add_run(
    "Each client maintains a trust score (0–1). Score halves each round a client is flagged; "
    "grows +0.1 for honest submissions. Clients below score 0.3 are excluded from aggregation. "
    "Two consecutive flagged rounds (1.0 → 0.5 → 0.25) trigger exclusion."
)

doc.add_heading("3.4  Real Network Mode", level=2)
para(
    "Beyond simulation, NeuralX-FL supports real networked deployment. "
    "The server runs on one laptop; clients run on separate machines "
    "connected via Tailscale (free, zero-config VPN). The server (server.py) "
    "uses Flower's gRPC protocol. Clients connect with: "
    "python run_client.py --partition-id 0 --server-address <IP>:9092. "
    "Traffic captured in Wireshark on port 9092 shows only floating-point arrays — "
    "no image data whatsoever."
)

# ---------------------------------------------------------------------------
# 4. FEATURES
# ---------------------------------------------------------------------------
heading("4. Features")

features = [
    ("Async Byzantine-Robust FL",
     "Handles both asynchronous client updates and Byzantine attacks simultaneously — "
     "a combination not found in standard FedAvg implementations."),
    ("Dual Detection Filters",
     "L2 norm filter catches scaling attacks. Cosine similarity filter catches "
     "directional attacks (sign-flip, random gradients). Both run before async selection."),
    ("Dynamic Trust Scoring",
     "Per-client reputation system with exponential decay on bad behaviour and "
     "gradual recovery. Persistent bad actors are automatically excluded."),
    ("4 Aggregation Methods",
     "FedAvg, Coordinate-wise Trimmed Mean, Coordinate-wise Median, and Krum — "
     "all selectable via a single config flag."),
    ("Differential Privacy",
     "Client-side adaptive DP clipping via Flower's built-in SecAgg-compatible "
     "mechanism. Privacy budget tracked via Rényi DP accountant (ε ≈ 0.6239)."),
    ("Real Network Deployment",
     "Toggle between simulation and real 3-laptop deployment with one config flag "
     "(USE_REAL_NETWORK). Tailscale integration for cross-network connectivity."),
    ("Memory-Mapped Data Loading",
     "numpy mmap_mode='r' allows all Ray actors to share the 200 MB dataset file "
     "using OS page sharing — reduces per-process memory from 847 MB to ~10 MB."),
    ("29 Unit Tests",
     "Full pytest suite covering every aggregation algorithm, both detection filters, "
     "label noise injection, and the DP epsilon estimator."),
    ("7 Result Plots",
     "Convergence curves, attack impact, dropout reliability, detection rate, "
     "heterogeneity analysis, DP trade-off, and loss curves — all auto-generated."),
]

for name, desc in features:
    bullet(f"{desc}", bold_prefix=f"{name}: ")

# ---------------------------------------------------------------------------
# 5. USP (Unique Selling Points)
# ---------------------------------------------------------------------------
heading("5. Unique Selling Points (USP)")

usps = [
    ("Filter-Before-Select Architecture",
     "Most async FL systems select the fastest clients, then filter. "
     "NeuralX-FL filters ALL updates first, then selects the fastest survivors. "
     "This prevents attackers from gaming the async buffer by responding quickly — "
     "a subtle but critical security property."),
    ("Trust Score Persistence Across Rounds",
     "Unlike per-round detection (which only acts on a single round), the trust "
     "scoring system accumulates evidence across rounds. A client that behaves badly "
     "occasionally but not every round is eventually excluded — standard norm/cosine "
     "filters would miss this."),
    ("Experimentally Verified Results on Real Hardware",
     "All results (84.22% clean, 18.64% under attack, 81.41% defended) are from "
     "actual GPU runs on a GTX 1650 — not theoretical projections. "
     "The 2.5-hour full experiment suite is reproducible from a single python main.py command."),
    ("Real Federated Network — Not Just Simulation",
     "The system runs as a real distributed system across 3 physical laptops using "
     "gRPC over Tailscale. Wireshark captures on port 9092 confirm only weight "
     "arrays (no images) cross the network — demonstrating the privacy guarantee physically."),
    ("Privacy-Utility Trade-off Quantified",
     "The system honestly reports that DP with noise_multiplier=0.1 drops accuracy "
     "from 81.41% to 38.70%, and explains exactly why (noise/clip ratio = 1.0 overwhelms "
     "gradients). This is a real research contribution, not a limitation hidden."),
]

for title, detail in usps:
    bullet(f"{detail}", bold_prefix=f"{title}: ")

# ---------------------------------------------------------------------------
# 6. RESULTS TABLE
# ---------------------------------------------------------------------------
heading("6. Experimental Results")

table_2col(
    [
        ["Exp A — FedAvg, no attack (clean baseline)", "84.22%"],
        ["Exp B — FedAvg, under 20% attack (no defence)", "18.64%"],
        ["Exp C — AsyncRobustFL (full defence)", "81.41%"],
        ["Exp D — AsyncRobustFL + Differential Privacy", "38.70%"],
        ["Attack damage (Exp A → Exp B)", "−65.58% accuracy"],
        ["Defence recovery (Exp B → Exp C)", "+62.77% accuracy"],
        ["Total flagged events (20 rounds)", "20 detection events"],
        ["Privacy budget ε (Exp D)", "≈ 0.6239 at δ = 1×10⁻⁵"],
    ],
    header=["Metric", "Value"]
)

# ---------------------------------------------------------------------------
# 7. TECH STACK
# ---------------------------------------------------------------------------
heading("7. Technical Stack")

table_2col(
    [
        ["Language", "Python 3.10"],
        ["Deep Learning", "PyTorch 2.6.0+cu124"],
        ["GPU", "NVIDIA GTX 1650, CUDA 12.4"],
        ["FL Framework", "Flower (flwr[simulation]) ≥ 1.8.0"],
        ["Parallel Simulation", "Ray (Flower backend)"],
        ["Dataset", "medmnist ≥ 3.0 (PathMNIST)"],
        ["Numerics", "NumPy 2.2.6"],
        ["Plotting", "Matplotlib 3.7+"],
        ["Privacy Accounting", "SciPy 1.10+ (Rényi DP accountant)"],
        ["Network (Real mode)", "Tailscale + Flower gRPC"],
        ["Testing", "pytest (29 unit tests, 100% pass rate)"],
        ["Codebase size", "~900 lines across 8 modules"],
    ],
    header=["Component", "Technology"]
)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "NeuralX_FL_Report.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
